from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from typing import Dict, Optional, List
from ._types import Paragraph, _Cell, Document

ROW_START, ROW_END = 2, 33
FIRST_HALF_COL_START, FIRST_HALF_COL_END = 0, 3
SECOND_HALF_COL_START, SECOND_HALF_COL_END = 4, 7


class TableStacker:
    """
    The TableStacker object merges the table of two columns into one as in template.docx, and structures the merged column into a dict.
    """

    def __init__(self, doc: Document) -> None:
        self.table = doc.tables[0]
        self.table.style.font.name = "BiauKai"

    def render_dict(
        self, row_start: int = ROW_START, row_end: int = ROW_END
    ) -> Dict[int, Dict]:

        first_half = []
        secodn_half = []

        for i in range(row_start, row_end + 1):
            row = self.table.row_cells(i)
            first_half += row[FIRST_HALF_COL_START : FIRST_HALF_COL_END + 1]
            secodn_half += row[SECOND_HALF_COL_START : SECOND_HALF_COL_END + 1]

        all_ = first_half + secodn_half

        indexed_cell = {}

        # A cell unit contains 8 element: first 4 are row1, second 4, row2
        for i in range(len(all_) // 8):
            cell = all_[8 * i : 8 * i + 8]
            indexed_cell[i] = {
                "row1": cell[:4],
                "row2": cell[4:],
            }
        return indexed_cell


class Cell:
    """
    The Cell object is a wrapper class for docx.table._Cell, handling the modification of cells in a Grid object
    """

    def __init__(self, cell: _Cell) -> None:
        self._cell = cell

    def apply_basic_format(self, paragraph: Paragraph) -> None:
        paragraph.style.font.size = Pt(14)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_text(self, text: Optional[str]) -> None:
        if text is None:
            return 
        paragraph_in_cell = self._cell.paragraphs[0]
        paragraph_in_cell.text = text
        self.apply_basic_format(paragraph_in_cell)

    def add_signature(self, signature_path: Optional[str]) -> None:
        if signature_path is None:
            return
        paragraph_in_cell = self._cell.paragraphs[0]
        run = paragraph_in_cell.add_run()
        run.add_tab()
        run.add_picture(signature_path, width=Cm(2.8), height=Cm(1.15))

    def __repr__(self):
        return self._cell.paragraphs[0].text

    __str__ = __repr__


class Grid:
    """
    The Grid object contains 6 Cell objects: date, work_hours, signature_start, signature_end, start_time, and end_time.
    They can be modified through the methods of the Cell object.
    """

    def __init__(self, grid: Dict[str, List[_Cell]]) -> None:
        self._row1 = grid["row1"]
        self._row2 = grid["row2"]

    @property
    def date(self) -> Cell:
        return Cell(self._row1[0])

    @property
    def work_hours(self) -> Cell:
        return Cell(self._row1[3])

    @property
    def signature_start(self) -> Cell:
        return Cell(self._row1[1])

    @property
    def signature_end(self) -> Cell:
        return Cell(self._row1[2])

    @property
    def start_time(self) -> Cell:
        return Cell(self._row2[1])

    @property
    def end_time(self) -> Cell:
        return Cell(self._row2[2])

    def fill_data(self, **kwargs) -> None:
        date = kwargs.get("date")
        start_time = kwargs.get("start_time")
        end_time = kwargs.get("end_time")
        signature_start = kwargs.get("signature_start")
        signature_end = kwargs.get("signature_end")
        work_hours = kwargs.get("work_hours")
        if not all(
            [date, start_time, end_time, signature_end, signature_start, work_hours]
        ):
            raise ValueError(f"None in: {kwargs}")
        self.date.add_text(date)
        self.start_time.add_text(start_time)
        self.end_time.add_text(end_time)
        self.signature_end.add_signature(signature_end)
        self.signature_start.add_signature(signature_start)
        self.work_hours.add_text(work_hours)


class CellDataGenerator:
    """
    The CellDataGenerator object generates all the required data to fill in a Grid object, based on the given year,
    month, start_time, work_hours, work_day and the signature path.
    """

    def __init__(
        self,
        year: int,
        month: int,
        start_time: str,
        work_hours: int,
        work_day: int,
        signature_path: str,
    ) -> None:
        self._start_year_month = f"{year}/{month}"
        self._end_year_month = f"{year}/{month + 1}"
        self._start_time = pd.Timestamp(start_time)
        self._end_time = pd.Timestamp(start_time) + pd.to_timedelta(
            work_hours, unit="H"
        )
        self._work_hours = work_hours
        self._work_day = int(work_day)
        self._signature_path = signature_path

    def render_dict(self, orient: str = "index") -> Dict[int, Dict[str, str]]:
        """
        Args:
            orient: Determining the mapping between keys and values, and passed to pandas.DataFrame.to_dict, .
                    More info at: https://pandas.pydata.org/pandas-docs/stable/reference/api/pandas.DataFrame.to_dict.html

        Returns:
            A dict: {
                        index0: {column0: value, column1: value, ...},
                        index1: {column0: value, column1: value, ...},
                        ...
                    }

        Raises:
            AssertionError: raised when work_days is, say, 25, but there are only, say, 22 or 23 bussiness days
        """

        time_table = pd.DataFrame()

        time_table["date"] = (
            pd.date_range(
                start=self._start_year_month,
                end=self._end_year_month,
                closed="left",
                freq="B",
            )
        ).strftime("%m/%d")

        assert (
            len(time_table["date"]) >= self._work_day
        ), f"Only {len(time_table)} bussiness days, but got {self._work_day} work days"

        time_table["start_time"] = self._start_time.strftime("%H:%M")
        time_table["end_time"] = self._end_time.strftime("%H:%M")
        time_table["work_hours"] = str(self._work_hours)
        time_table["signature_start"] = self._signature_path
        time_table["signature_end"] = self._signature_path

        time_table = (
            time_table.sample(n=self._work_day)
            .sort_values("date")
            .reset_index(drop=True)
        )

        indexed_dict: Dict[int, Dict[str, str]] = time_table.to_dict(orient)

        return indexed_dict