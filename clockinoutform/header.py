from docx.text.run import Run


class RunText:
    def __init__(self, run):
        if not isinstance(run, Run):
            raise TypeError(f"{run}: {type(run).__name__} not an instance of Run")
        self._run = run

    def add_text(self, text, padding_length=2):
        self._run.text = self._pad(text, padding_length, " ")
        return self._run.text

    def _pad(self, text, length, padding):
        length_rjust = len(text) + length
        length_ljust = len(text) + length * 2
        return text.rjust(length_rjust, padding).ljust(length_ljust, padding)


class BasicInfo:
    def __init__(self, doc):
        self.doc = doc

    @property
    def year(self):
        addr = self.doc.paragraphs[0].runs[1]
        return addr, addr.text

    @year.setter
    def year(self, value):
        return RunText(self.year[0]).add_text(value)

    @property
    def month(self):
        addr = self.doc.paragraphs[0].runs[3]
        return addr, addr.text.zfill(2)

    @month.setter
    def month(self, value):
        value = str(value).zfill(2)
        return RunText(self.month[0]).add_text(value)

    @property
    def department(self):
        addr = self.doc.paragraphs[1].runs[1]
        return addr, addr.text

    @department.setter
    def department(self, value):
        return RunText(self.department[0]).add_text(value)

    @property
    def name(self):
        addr = self.doc.paragraphs[1].runs[6]
        return addr, addr.text

    @name.setter
    def name(self, value):
        return RunText(self.name[0]).add_text(value)

    @property
    def expected(self):
        addr = self.doc.paragraphs[1].runs[9]
        return addr, addr.text

    @expected.setter
    def expected(self, value):
        return RunText(self.expected[0]).add_text(value)

    @property
    def actual(self):
        addr = self.doc.paragraphs[1].runs[13]
        return addr, addr.text

    @actual.setter
    def actual(self, value):
        return RunText(self.actual[0]).add_text(value)
